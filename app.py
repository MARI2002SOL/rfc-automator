import io
import re
import zipfile
from datetime import datetime
import docx
from google import genai
import streamlit as st
from streamlit_extras.let_it_rain import rain
from streamlit_option_menu import option_menu

# ==========================================
# CONFIGURACIÓN DE GEMINI API (SDK MODERNO)
# ==========================================
client = None
try:
    gemini_api_key = st.secrets["GEMINI_API_KEY"]
    client = genai.Client(api_key=gemini_api_key)
except Exception:
    st.warning(
        "⚠️ No se detectó 'GEMINI_API_KEY' en st.secrets. Agrega la clave para poder usar la IA."
    )


# ==========================================
# FUNCIONES HELPER - PARTE 1 (DOCX & TXT)
# ==========================================
def reemplazar_texto_en_parrafo(parrafo, mapa_reemplazos):
    for buscar, reemplazo in mapa_reemplazos.items():
        if buscar in parrafo.text:
            parrafo.text = parrafo.text.replace(buscar, reemplazo)


def reemplazar_manteniendo_formato(doc, mapa_reemplazos):
    """Reemplaza los marcadores en párrafos y tablas del documento Word."""
    for parrafo in doc.paragraphs:
        reemplazar_texto_en_parrafo(parrafo, mapa_reemplazos)
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for parrafo in celda.paragraphs:
                    reemplazar_texto_en_parrafo(parrafo, mapa_reemplazos)


def obtener_valor_exacto(doc, etiqueta_buscada):
    """Busca el valor a la derecha o abajo de una etiqueta dada en las tablas."""
    etiqueta_clean = etiqueta_buscada.strip().lower()
    for tabla in doc.tables:
        for i_fila, fila in enumerate(tabla.rows):
            for i_celda, celda in enumerate(fila.cells):
                texto_celda = celda.text.strip().lower()
                if texto_celda == etiqueta_clean:
                    if i_celda + 1 < len(fila.cells):
                        val_derecha = fila.cells[i_celda + 1].text.strip()
                        if val_derecha and val_derecha.lower() != texto_celda:
                            return val_derecha
                    if i_fila + 1 < len(tabla.rows):
                        val_abajo = (
                            tabla.rows[i_fila + 1].cells[i_celda].text.strip()
                        )
                        if val_abajo:
                            return val_abajo
    return ""


def obtener_tarea_plan(doc, nombre_plan):
    """Obtiene el texto de los planes de ejecución o reversión en las tablas."""
    nombre_plan_clean = nombre_plan.strip().lower()
    for tabla in doc.tables:
        for fila in tabla.rows:
            celdas_texto = []
            for celda in fila.cells:
                txt = celda.text.strip()
                if not celdas_texto or celdas_texto[-1] != txt:
                    celdas_texto.append(txt)
            for idx, txt in enumerate(celdas_texto):
                if nombre_plan_clean in txt.lower():
                    for posterior in celdas_texto[idx + 1 :]:
                        if posterior and posterior.lower() != txt.lower():
                            return posterior
    return ""


# ==========================================
# GENERACIÓN DE SQL
# ==========================================
def generar_queries_sql_con_gemini(
    texto_sunat,
    tipo_operacion="INSERT",
    ticket="123456",
    estado_previo="INACTIVO",
):
    prompt_sistema = f"""
    Eres un DBA experto en SQL Server para sistemas peruanos.
    Tu trabajo es procesar texto copiado de consultas RUC de SUNAT y generar exclusivamente scripts SQL.

    ESTRUCTURA DE LA TABLA `sunat_contribuyente`:
    - numero_ruc (varchar)
    - razon_social (varchar)
    - estado (varchar) -> Solo el estado (ej: 'ACTIVO', 'BAJA DE OFICIO')
    - condicion_domicilio (varchar) -> Solo la condición (ej: 'HABIDO', 'NO HABIDO')
    - ubigeo (varchar de 6 dígitos) -> CONVIERTE la cadena "DPTO - PROV - DIST" al código oficial de Ubigeo INE de 6 dígitos (ej: '140101', '150110', '140105').
    - tipo_via (varchar) -> ej: 'CAL.', 'AV.', 'JR.' (si aplica)
    - nombre_via (varchar) -> Nombre de la vía (si aplica)
    - numero (varchar) -> Número de puerta (si aplica)
    - interior (varchar) -> (si aplica)
    - lote (varchar) -> (si aplica)
    - departamento (varchar) -> (si aplica)
    - manzana (varchar) -> (si aplica)
    - kilometro (varchar) -> (si aplica)
    - tipo_zona (varchar) -> ej: 'URB.', 'A.H.' (si aplica)
    - codigo_zona (varchar) -> Nombre de la zona (si aplica)
    - estado_fila (varchar) -> 'I' si es INSERT, 'U' si es UPDATE.

    REGLAS DE OPERACIÓN:
    1. Responde ÚNICAMENTE con dos bloques separados por el delimitador "===ROLLBACK_SEPARADOR===":
       - Primero la consulta de PASE A PROD.
       - Luego el delimitador.
       - Finalmente la consulta de ROLLBACK.
    2. El encabezado del PASE debe iniciar con `--{ticket}`.
    3. La base de datos es `labroe_new`. Usa la sintaxis:
       use labroe_new
       go
    4. Omite los campos que sean NULL o no existan en la información dada.
    5. NO incluyas bloques de código Markdown (```sql), devuelve solo texto plano.
    6. Para ROLLBACK de UPDATE e INSERT: usa `update sunat_contribuyente set estado = '{estado_previo}', fecha_actualizacion = getdate() where numero_ruc = '...';`
    """

    prompt_usuario = f"""
    TIPO OPERACIÓN: {tipo_operacion}
    TEXTO SUNAT:
    {texto_sunat}
    """

    response = client.models.generate_content(
        model="gemini-3.6-flash",
        contents=prompt_usuario,
        config={"system_instruction": prompt_sistema},
    )

    partes = response.text.split("===ROLLBACK_SEPARADOR===")
    q_prod = partes[0].strip() if len(partes) > 0 else response.text.strip()
    q_roll = partes[1].strip() if len(partes) > 1 else ""

    q_prod = re.sub(r"^```sql\n?|^```\n?", "", q_prod, flags=re.MULTILINE)
    q_prod = re.sub(r"\n?```$", "", q_prod, flags=re.MULTILINE)

    q_roll = re.sub(r"^```sql\n?|^```\n?", "", q_roll, flags=re.MULTILINE)
    q_roll = re.sub(r"\n?```$", "", q_roll, flags=re.MULTILINE)

    return q_prod, q_roll


def generar_query_actualizacion_usuario(dni, id_usuario):
    query = f"""use labroe_new
go

-- COLOCAR DNI ACTUAL DEL USUARIO
update x set x.usuario = '{dni}', x.contrasena_temporal = '{dni}'
from RoeRestService.usuario_web x
-- COLOCAR ID DEL USUARIO
where x.id_usuario = {id_usuario}
go"""
    return query


def generar_queries_homologacion(lista_analisis, ticket_num):
    """Genera las queries de PASE y ROLLBACK para la Homologación de Exámenes."""
    lines_pase = ["USE labroe_new", "GO", f"-- TICKET: {ticket_num}", ""]
    lines_roll = ["USE labroe_new", "GO", f"-- ROLLBACK TICKET: {ticket_num}", ""]

    for idx, item in enumerate(lista_analisis, 1):
        if not item.get("es_update", False):
            # CASO INSERT
            roe = item.get("codana_roe", "").strip()
            seq = item.get("codana_sequence", "").strip()
            if roe and seq:
                lines_pase.append(f"-- CASO {idx}: SOLO INSERT")
                lines_pase.append(
                    f"insert into interface values (2, '{roe}', 53, '{seq}');\n"
                )

                lines_roll.append(f"-- ROLLBACK CASO {idx}: DELETE")
                lines_roll.append(
                    f"delete from interface where id_codigo_sistema = 2 and id_codigo_interno = '{roe}' and id_codigo_externo = '{seq}';\n"
                )
        else:
            # CASO UPDATE
            cod_interno = item.get("codigo_silc", "").strip()
            cod_ext_antiguo = item.get("codigo_externo_antiguo", "").strip()
            cod_ext_nuevo = item.get("codigo_externo_nuevo", "").strip()

            if cod_interno and cod_ext_antiguo and cod_ext_nuevo:
                # --- QUERY PASE A PRODUCCIÓN ---
                lines_pase.append(f"-- CASO {idx}: UPDATE HOMOLOGACIÓN")
                lines_pase.append(
                    "UPDATE interface\n"
                    f"SET id_codigo_externo = '{cod_ext_nuevo}'\n"
                    f"WHERE id_codigo_interno = '{cod_interno}'\n"
                    f"  AND id_codigo_externo = '{cod_ext_antiguo}';\n"
                )

            # --- QUERY ROLLBACK ---
                lines_roll.append(f"-- ROLLBACK CASO {idx}: REVERSIÓN UPDATE")
                lines_roll.append(
                    "UPDATE interface\n"
                    f"SET id_codigo_externo = '{cod_ext_antiguo}'\n"
                    f"WHERE id_codigo_interno = '{cod_interno}'\n"
                    f"  AND id_codigo_externo = '{cod_ext_nuevo}';\n"
                )

    lines_pase.append("GO")
    lines_roll.append("GO")

    return "\n".join(lines_pase), "\n".join(lines_roll)


# ==========================================
# INTERFAZ STREAMLIT
# ==========================================
st.set_page_config(page_title="Gestor RUC & RFC", page_icon="📄", layout="wide")

with st.sidebar:
    opcion = option_menu(
        menu_title="Menú Principal",
        options=[
            "Actualizar/Registrar RUC",
            "Actualizar credenciales",
            "Homologación de exámenes",
        ],
        icons=["house", "folder", "envelope"],
        menu_icon="cast",
        default_index=0,
    )

# ---------------------------------------------------------
# OPCIÓN 1: ACTUALIZAR / REGISTRAR RUC
# ---------------------------------------------------------
if opcion == "Actualizar/Registrar RUC":
    st.title("📄 Sistema de Gestión Unificado RUC / RFC")

    if "resultado_procesado_ruc" not in st.session_state:
        st.session_state.resultado_procesado_ruc = None

    st.header("1. Carga de Archivo y Datos")
    col_left, col_right = st.columns(2)

    with col_left:
        uploaded_file = st.file_uploader(
            "Sube la Solicitud de Cambio (.docx)",
            type=["docx"],
            key="file_ruc",
        )
        ticket_num = st.text_input(
            "Ingresa el N° de Ticket:", placeholder="Ej: 12776", key="ticket_ruc"
        )

    with col_right:
        texto_ruc = st.text_area(
            "Pega la Consulta RUC copiada de [SUNAT](https://e-consultaruc.sunat.gob.pe/cl-ti-itmrconsruc/FrameCriterioBusquedaWeb.jsp):",
            height=180,
            key="txt_ruc",
        )
        estado_previo = st.text_input(
            "Estado anterior para Rollback:", value="INACTIVO", key="prev_ruc"
        )

    tipo_operacion = "INSERT"
    if uploaded_file is not None:
        nombre_archivo = uploaded_file.name.lower()
        if "actualizar" in nombre_archivo:
            tipo_operacion = "UPDATE"
        elif "registrar" in nombre_archivo or "registro" in nombre_archivo:
            tipo_operacion = "INSERT"

    st.info(
        f"📌 **Tipo de Operación detectado automáticamente:** `{tipo_operacion}`"
    )

    if st.button("🚀 Procesar Todo y Generar Archivos", key="btn_procesar_ruc"):
        rain(
            emoji="📄", font_size=54, falling_speed=5, animation_length=1
        )
        if not uploaded_file or not ticket_num or not texto_ruc:
            st.error("Por favor completa todos los campos requeridos.")
        elif not client:
            st.error("No se ha configurado 'GEMINI_API_KEY' en los secretos.")
        else:
            with st.spinner(
                "Procesando documento Word y generando queries SQL..."
            ):
                try:
                    doc = docx.Document(uploaded_file)
                    fecha_actual = datetime.now().strftime("%d/%m/%Y")
                    reemplazos = {
                        "[FECHA DE HOY]": fecha_actual,
                        "TICKETNUM": ticket_num,
                    }
                    reemplazar_manteniendo_formato(doc, reemplazos)

                    detalle_cambio = obtener_valor_exacto(doc, "Descripción")
                    plan_ejecucion = obtener_tarea_plan(doc, "ejecución")
                    plan_reversion = obtener_tarea_plan(
                        doc, "PLAN DE REVERSIÓN"
                    )
                    descripcion_cambio = obtener_valor_exacto(
                        doc, "MOTIVO DEL CAMBIO"
                    )
                    analista_responsable = obtener_valor_exacto(doc, "Nombre")

                    buffer_docx = io.BytesIO()
                    doc.save(buffer_docx)
                    buffer_docx.seek(0)

                    q_prod, q_rollback = generar_queries_sql_con_gemini(
                        texto_ruc, tipo_operacion, ticket_num, estado_previo
                    )

                    nombre_docx_out = f"{uploaded_file.name.replace('.docx', '')}_{ticket_num}.docx"
                    if tipo_operacion == "INSERT":
                        nom_prod = (
                            f"query_RUC_registro - PASE A PROD ({ticket_num}).sql"
                        )
                        nom_roll = (
                            f"query_RUC_registro - ROLLBACK ({ticket_num}).sql"
                        )
                    else:
                        nom_prod = f"query_RUC_actualizar - PASE A PROD ({ticket_num}).sql"
                        nom_roll = (
                            f"query_RUC_actualizar - ROLLBACK ({ticket_num}).sql"
                        )

                    zip_buffer = io.BytesIO()
                    with zipfile.ZipFile(
                        zip_buffer, "w", zipfile.ZIP_DEFLATED
                    ) as zip_file:
                        zip_file.writestr(
                            nombre_docx_out, buffer_docx.getvalue()
                        )
                        zip_file.writestr(nom_prod, q_prod)
                        zip_file.writestr(nom_roll, q_rollback)

                    zip_buffer.seek(0)

                    st.session_state.resultado_procesado_ruc = {
                        "q_prod": q_prod,
                        "q_rollback": q_rollback,
                        "zip_data": zip_buffer.getvalue(),
                        "zip_name": f"RFC_RUC_{ticket_num}.zip",
                        "detalle_cambio": detalle_cambio,
                        "plan_ejecucion": plan_ejecucion,
                        "plan_reversion": plan_reversion,
                        "descripcion_cambio": descripcion_cambio,
                        "analista_responsable": analista_responsable,
                    }
                except Exception as e:
                    st.error(f"Error durante el procesamiento: {e}")

    if st.session_state.resultado_procesado_ruc:
        st.divider()
        st.header("2. Vista Previa de Scripts SQL Generados")
        res = st.session_state.resultado_procesado_ruc

        col_p1, col_p2 = st.columns(2)
        with col_p1:
            st.subheader("Pase a Producción (PASE)")
            st.code(res["q_prod"], language="sql")
        with col_p2:
            st.subheader("Rollback (ROLLBACK)")
            st.code(res["q_rollback"], language="sql")

        st.divider()
        st.header("3. Confirmación y Descarga Unificada")
        st.download_button(
            label="📦 DESCARGAR TODOS LOS ARCHIVOS (.ZIP)",
            data=res["zip_data"],
            file_name=res["zip_name"],
            mime="application/zip",
            type="primary",
            key="dl_ruc",
        )

        st.divider()
        st.header("Información adicional")
        col_f1_1, col_f1_2 = st.columns(2)
        with col_f1_1:
            st.subheader("Detalle del Cambio/Despliegue")
            texto_detalle = (
                f"{res['detalle_cambio']}\n\n"
                f"PLAN DE EJECUCIÓN\n{res['plan_ejecucion']}\n\n"
                f"PLAN DE REVERSIÓN (Roll-back)\n{res['plan_reversion']}"
            )
            st.code(texto_detalle, language="text")
        with col_f1_2:
            st.subheader("Descripción del cambio")
            st.code(res["descripcion_cambio"], language="text")

        col_f2_1, col_f2_2 = st.columns(2)
        with col_f2_1:
            st.subheader(
                "Analista/Especialista responsable del despliegue del cambio"
            )
            st.code(res["analista_responsable"], language="text")
        with col_f2_2:
            st.subheader("¿Existe Riesgo?")
            st.code("NINGUNO", language="text")


# ---------------------------------------------------------
# OPCIÓN 2: ACTUALIZAR CREDENCIALES
# ---------------------------------------------------------
elif opcion == "Actualizar credenciales":
    st.title("🪪 Sistema de Gestión Unificado RUC / RFC")

    if "resultado_procesado_cred" not in st.session_state:
        st.session_state.resultado_procesado_cred = None

    st.header("1. Carga de Archivo y Datos")
    col_left, col_right = st.columns(2)

    with col_left:
        uploaded_file = st.file_uploader(
            "Sube la Solicitud de Cambio (.docx)",
            type=["docx"],
            key="file_cred",
        )
        ticket_num = st.text_input(
            "Ingresa el N° de Ticket:", placeholder="Ej: 12776", key="ticket_cred"
        )

    with col_right:
        dni_antiguo = st.text_input(
            "¿Cuál es el número de DNI antiguo?",
            placeholder="Ej: 11111111",
            key="dni_ant",
        )
        dni_nuevo = st.text_input(
            "¿Cuál es el número de DNI nuevo?",
            placeholder="Ej: 22222222",
            key="dni_nue",
        )
        user_id = st.text_input(
            "¿Cuál es el ID del usuario?",
            placeholder="Ej: 3333333",
            key="uid_cred",
        )

    if uploaded_file is not None:
        if "solicitud" not in uploaded_file.name.lower():
            st.error("⚠️ Por favor, suba un archivo correcto.")
        else:
            st.success("¡Archivo cargado correctamente!")

    if st.button(
        "🚀 Procesar Todo y Generar Archivos", key="btn_procesar_credenciales"
    ):
        rain(
            emoji="🪪", font_size=54, falling_speed=5, animation_length=1
        )
        if (
            not uploaded_file
            or not ticket_num
            or not dni_antiguo
            or not dni_nuevo
            or not user_id
        ):
            st.error("Por favor completa todos los campos requeridos.")
        else:
            with st.spinner(
                "Procesando documento Word y generando queries SQL..."
            ):
                try:
                    doc = docx.Document(uploaded_file)
                    fecha_actual = datetime.now().strftime("%d/%m/%Y")
                    reemplazos = {
                        "[FECHA DE HOY]": fecha_actual,
                        "TICKETNUM": ticket_num,
                    }
                    reemplazar_manteniendo_formato(doc, reemplazos)

                    detalle_cambio = obtener_valor_exacto(doc, "Descripción")
                    plan_ejecucion = obtener_tarea_plan(doc, "ejecución")
                    plan_reversion = obtener_tarea_plan(
                        doc, "PLAN DE REVERSIÓN"
                    )
                    descripcion_cambio = obtener_valor_exacto(
                        doc, "MOTIVO DEL CAMBIO"
                    )
                    analista_responsable = obtener_valor_exacto(doc, "Nombre")

                    buffer_docx = io.BytesIO()
                    doc.save(buffer_docx)
                    buffer_docx.seek(0)

                    q_prod = generar_query_actualizacion_usuario(
                        dni_nuevo, user_id
                    )
                    q_rollback = generar_query_actualizacion_usuario(
                        dni_antiguo, user_id
                    )

                    nombre_docx_out = f"{uploaded_file.name.replace('.docx', '')}_{ticket_num}.docx"
                    nom_prod = f"({ticket_num})_paseprod.sql"
                    nom_roll = f"({ticket_num})_rollback.sql"

                    zip_buffer = io.BytesIO()
                    with zipfile.ZipFile(
                        zip_buffer, "w", zipfile.ZIP_DEFLATED
                    ) as zip_file:
                        zip_file.writestr(
                            nombre_docx_out, buffer_docx.getvalue()
                        )
                        zip_file.writestr(nom_prod, q_prod)
                        zip_file.writestr(nom_roll, q_rollback)

                    zip_buffer.seek(0)

                    st.session_state.resultado_procesado_cred = {
                        "q_prod": q_prod,
                        "q_rollback": q_rollback,
                        "zip_data": zip_buffer.getvalue(),
                        "zip_name": f"RFC_Credenciales_{ticket_num}.zip",
                        "detalle_cambio": detalle_cambio,
                        "plan_ejecucion": plan_ejecucion,
                        "plan_reversion": plan_reversion,
                        "descripcion_cambio": descripcion_cambio,
                        "analista_responsable": analista_responsable,
                    }
                except Exception as e:
                    st.error(f"Error durante el procesamiento: {e}")

    if st.session_state.resultado_procesado_cred:
        st.divider()
        st.header("2. Vista Previa de Scripts SQL Generados")
        res = st.session_state.resultado_procesado_cred

        col_p1, col_p2 = st.columns(2)
        with col_p1:
            st.subheader("Pase a Producción (PASE)")
            st.code(res["q_prod"], language="sql")
        with col_p2:
            st.subheader("Rollback (ROLLBACK)")
            st.code(res["q_rollback"], language="sql")

        st.divider()
        st.header("3. Confirmación y Descarga Unificada")
        st.download_button(
            label="📦 DESCARGAR TODOS LOS ARCHIVOS (.ZIP)",
            data=res["zip_data"],
            file_name=res["zip_name"],
            mime="application/zip",
            type="primary",
            key="dl_cred",
        )

        st.divider()
        st.header("Información adicional")
        col_f1_1, col_f1_2 = st.columns(2)
        with col_f1_1:
            st.subheader("Detalle del Cambio/Despliegue")
            texto_detalle = (
                f"{res['detalle_cambio']}\n\n"
                f"PLAN DE EJECUCIÓN\n{res['plan_ejecucion']}\n\n"
                f"PLAN DE REVERSIÓN (Roll-back)\n{res['plan_reversion']}"
            )
            st.code(texto_detalle, language="text")
        with col_f1_2:
            st.subheader("Descripción del cambio")
            st.code(res["descripcion_cambio"], language="text")

        col_f2_1, col_f2_2 = st.columns(2)
        with col_f2_1:
            st.subheader(
                "Analista/Especialista responsable del despliegue del cambio"
            )
            st.code(res["analista_responsable"], language="text")
        with col_f2_2:
            st.subheader("¿Existe Riesgo?")
            st.code("NINGUNO", language="text")


# ---------------------------------------------------------
# OPCIÓN 3: HOMOLOGACIÓN DE EXÁMENES
# ---------------------------------------------------------
# ---------------------------------------------------------
# OPCIÓN 3: HOMOLOGACIÓN DE EXÁMENES
# ---------------------------------------------------------
elif opcion == "Homologación de exámenes":
    st.title("📄 Sistema de Gestión Unificado RFC - Homologación")

    if "resultado_procesado_homo" not in st.session_state:
        st.session_state.resultado_procesado_homo = None

    if "lista_analisis" not in st.session_state:
        st.session_state.lista_analisis = [
            {
                "es_update": False,
                "codana_roe": "",
                "codana_sequence": "",
                "codigo_externo_nuevo": "",
                "codigo_silc": "",
                "codigo_externo_antiguo": "",
            }
        ]

    # CALLBACK: Sincroniza el cambio del checkbox en tiempo real antes de re-renderizar la UI
    def cambiar_estado_update(posicion):
        key = f"chk_update_{posicion}"
        st.session_state.lista_analisis[posicion]["es_update"] = st.session_state[key]

    st.header("1. Carga de Archivo y Datos")

    col_left, col_right = st.columns([1, 2])

    # --- COLUMNA IZQUIERDA: CONFIGURACIÓN GENERAL Y RESUMEN ---
    with col_left:
        with st.container(border=True):
            st.subheader("📋 Datos del Ticket")
            uploaded_file = st.file_uploader(
                "Sube la Solicitud de Cambio (.docx)",
                type=["docx"],
                key="file_homologacion",
            )
            ticket_num = st.text_input(
                "Ingresa el N° de Ticket:",
                placeholder="Ej: 12776",
                key="ticket_homologacion",
            )

        # Métricas calculadas sobre la fuente única de verdad (lista_analisis)
        total_items = len(st.session_state.lista_analisis)
        num_updates = sum(1 for item in st.session_state.lista_analisis if item["es_update"])
        num_inserts = total_items - num_updates

        with st.container(border=True):
            st.subheader("📊 Resumen a Procesar")
            m1, m2, m3 = st.columns(3)
            m1.metric("Total Análisis", total_items)
            m2.metric("N° Inserts", num_inserts)
            m3.metric("N° Updates", num_updates)
            if num_updates > 0:
                st.caption(f"⚡ *Incluye {num_updates} registro(s) tipo UPDATE*")

        with st.expander("💡 Guía de llenado", expanded=False):
            st.markdown(
                """
                * **INSERT (por defecto):** Ingresa `Cod. ROE` y `Cod. SEQUENCE`.
                * **UPDATE:** Marca la casilla e ingresa `Cod. Interno`, `SILC Antiguo` y `SILC Nuevo`.
                * Usa los botones **➕** y **🗑️** para ajustar el número de exámenes.
                """
            )

    # --- COLUMNA DERECHA: REGISTRO DINÁMICO DE ANÁLISIS ---
    with col_right:
        with st.container(border=True):
            st.subheader("🧪 Análisis a Homologar")

            col_btn1, col_btn2 = st.columns(2)
            with col_btn1:
                if st.button("➕ Agregar Análisis", key="btn_add_analisis", use_container_width=True):
                    st.session_state.lista_analisis.append(
                        {
                            "es_update": False,
                            "codana_roe": "",
                            "codana_sequence": "",
                            "codigo_externo_nuevo": "",
                            "codigo_silc": "",
                            "codigo_externo_antiguo": "",
                        }
                    )
                    st.rerun()

            with col_btn2:
                if (
                    st.button("🗑️ Quitar Último", key="btn_del_analisis", use_container_width=True)
                    and len(st.session_state.lista_analisis) > 1
                ):
                    st.session_state.lista_analisis.pop()
                    st.rerun()

            st.divider()

            for idx, item in enumerate(st.session_state.lista_analisis):
                # Estado actual sincronizado
                es_update_val = item["es_update"]
                tipo_tag = "UPDATE" if es_update_val else "INSERT"

                with st.expander(f"Análisis #{idx + 1} — [{tipo_tag}]", expanded=True):
                    st.checkbox(
                        "¿Es un UPDATE? (Marcar solo si requiere actualizar)",
                        value=es_update_val,
                        key=f"chk_update_{idx}",
                        on_change=cambiar_estado_update,
                        args=(idx,),
                    )

                    if not es_update_val:
                        c1, c2 = st.columns(2)
                        with c1:
                            item["codana_roe"] = st.text_input(
                                "Cod. ROE:",
                                value=item["codana_roe"],
                                key=f"roe_{idx}",
                                placeholder="Ej: PECYQ00",
                            )
                        with c2:
                            item["codana_sequence"] = st.text_input(
                                "Cod. SEQUENCE:",
                                value=item["codana_sequence"],
                                key=f"seq_{idx}",
                                placeholder="Ej: Z902400",
                            )
                    else:
                        c1, c2, c3 = st.columns(3)

                        with c1:
                            item["codigo_externo_nuevo"] = st.text_input(
                                "Cod. SEQUENCE nuevo:",
                                value=item["codigo_externo_nuevo"],
                                key=f"nue_{idx}",
                                placeholder="Ej: Z902400",
                            )


                        with c2:
                            item["codigo_silc"] = st.text_input(
                                "Código SILC:",
                                value=item["codigo_silc"],
                                key=f"int_{idx}",
                                placeholder="Ej: PXTL000",
                            )

                        with c3:

                            item["codigo_externo_antiguo"] = st.text_input(
                                "Cod. SEQUENCE antiguo:",
                                value=item["codigo_externo_antiguo"],
                                key=f"ant_{idx}",
                                placeholder="Ej: Z805300",
                            )
                            
    st.divider()
    if st.button(
        "🚀 Procesar Homologación y Generar Archivos",
        key="btn_procesar_homologacion",
        type="primary",
    ):
        rain(
            emoji="🧪", font_size=54, falling_speed=5, animation_length=1
        )
        if not uploaded_file or not ticket_num:
            st.error("Por favor adjunta la Solicitud (.docx) y el N° de Ticket.")
        else:
            with st.spinner(
                "Procesando documento Word y generando scripts SQL de homologación..."
            ):
                try:
                    doc = docx.Document(uploaded_file)
                    fecha_actual = datetime.now().strftime("%d/%m/%Y")
                    reemplazos = {
                        "[FECHA DE HOY]": fecha_actual,
                        "TICKETNUM": ticket_num,
                    }
                    reemplazar_manteniendo_formato(doc, reemplazos)

                    detalle_cambio = obtener_valor_exacto(doc, "Descripción")
                    plan_ejecucion = obtener_tarea_plan(doc, "ejecución")
                    plan_reversion = obtener_tarea_plan(
                        doc, "PLAN DE REVERSIÓN"
                    )
                    descripcion_cambio = obtener_valor_exacto(
                        doc, "MOTIVO DEL CAMBIO"
                    )
                    analista_responsable = obtener_valor_exacto(doc, "Nombre")

                    buffer_docx = io.BytesIO()
                    doc.save(buffer_docx)
                    buffer_docx.seek(0)

                    q_prod, q_rollback = generar_queries_homologacion(
                        st.session_state.lista_analisis, ticket_num
                    )

                    nombre_docx_out = f"{uploaded_file.name.replace('.docx', '')}_{ticket_num}.docx"
                    nom_prod = f"{ticket_num}_homologacion_paseprod.sql"
                    nom_roll = f"{ticket_num}_homologacion_rollback.sql"

                    zip_buffer = io.BytesIO()
                    with zipfile.ZipFile(
                        zip_buffer, "w", zipfile.ZIP_DEFLATED
                    ) as zip_file:
                        zip_file.writestr(
                            nombre_docx_out, buffer_docx.getvalue()
                        )
                        zip_file.writestr(nom_prod, q_prod)
                        zip_file.writestr(nom_roll, q_rollback)

                    zip_buffer.seek(0)

                    st.session_state.resultado_procesado_homo = {
                        "q_prod": q_prod,
                        "q_rollback": q_rollback,
                        "zip_data": zip_buffer.getvalue(),
                        "zip_name": f"RFC_Homologacion_{ticket_num}.zip",
                        "detalle_cambio": detalle_cambio,
                        "plan_ejecucion": plan_ejecucion,
                        "plan_reversion": plan_reversion,
                        "descripcion_cambio": descripcion_cambio,
                        "analista_responsable": analista_responsable,
                    }
                except Exception as e:
                    st.error(f"Error durante el procesamiento: {e}")

    if st.session_state.resultado_procesado_homo:
        st.divider()
        st.header("2. Vista Previa de Scripts SQL Generados")
        res = st.session_state.resultado_procesado_homo

        col_p1, col_p2 = st.columns(2)
        with col_p1:
            st.subheader("Pase a Producción (PASE)")
            st.code(res["q_prod"], language="sql")
        with col_p2:
            st.subheader("Rollback (ROLLBACK)")
            st.code(res["q_rollback"], language="sql")

        st.divider()
        st.header("3. Confirmación y Descarga Unificada")
        st.download_button(
            label="📦 DESCARGAR TODOS LOS ARCHIVOS (.ZIP)",
            data=res["zip_data"],
            file_name=res["zip_name"],
            mime="application/zip",
            type="primary",
            key="dl_homo",
        )

        st.divider()
        st.header("Información adicional")
        col_f1_1, col_f1_2 = st.columns(2)
        with col_f1_1:
            st.subheader("Detalle del Cambio/Despliegue")
            texto_detalle = (
                f"{res['detalle_cambio']}\n\n"
                f"PLAN DE EJECUCIÓN\n{res['plan_ejecucion']}\n\n"
                f"PLAN DE REVERSIÓN (Roll-back)\n{res['plan_reversion']}"
            )
            st.code(texto_detalle, language="text")
        with col_f1_2:
            st.subheader("Descripción del cambio")
            st.code(res["descripcion_cambio"], language="text")

        col_f2_1, col_f2_2 = st.columns(2)
        with col_f2_1:
            st.subheader(
                "Analista/Especialista responsable del despliegue del cambio"
            )
            st.code(res["analista_responsable"], language="text")
        with col_f2_2:
            st.subheader("¿Existe Riesgo?")
            st.code("NINGUNO", language="text")